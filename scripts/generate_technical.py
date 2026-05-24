#!/usr/bin/env python3
"""
技术方案生成器
功能：根据招标文件要求，自动生成技术方案章节内容

Usage:
    python generate_technical.py --tender-file tender.pdf --output technical_proposal.md
"""

import argparse
import json
import os
import re
import logging
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

# ============================================================
# 日志配置
# ============================================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# ============================================================
# 数据模型
# ============================================================

@dataclass
class Requirement:
    """招标文件中的技术要求"""
    clause_id: str        # 条款编号，如 "3.1.2"
    title: str            # 要求标题
    content: str          # 要求内容
    response: str         # 我方响应
    chapter: str          # 对应章节


@dataclass
class ProjectCase:
    """项目案例"""
    name: str
    customer: str
    contract_amount: float
    completion_date: str
    description: str
    tech_highlights: list[str] = field(default_factory=list)


# ============================================================
# PDF解析器
# ============================================================

class PDFParser:
    """PDF文件解析器"""

    @staticmethod
    def parse(file_path: str) -> str:
        """
        解析PDF文件，提取文本内容
        """
        import pdfplumber

        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text


# ============================================================
# DOCX解析器
# ============================================================

class DOCXParser:
    """Word文档解析器"""

    @staticmethod
    def parse(file_path: str) -> str:
        """
        解析Word文档，提取文本内容
        """
        from docx import Document

        text = ""
        doc = Document(file_path)

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text += row_text + "\n"
            text += "\n"

        return text


# ============================================================
# 招标文件解析器
# ============================================================

class TenderParser:
    """招标文件解析器，支持PDF和Word"""

    # 技术要求关键词
    QUALIFICATION_KEYWORDS = [
        "资格", "资质", "投标人", "响应", "承诺", "授权",
        "营业执照", "法人", "委托人", "证明材料"
    ]

    TECH_KEYWORDS = [
        "技术", "规格", "要求", "功能", "性能", "指标",
        "参数", "系统", "模块", "方案", "设计", "实施",
        "工期", "进度", "人员", "设备", "材料"
    ]

    EVALUATION_KEYWORDS = [
        "评分", "评审", "得分", "权重", "分值", "评标",
        "技术标", "商务标", "价格"
    ]

    def __init__(self, tender_file: str):
        self.tender_file = tender_file
        self.raw_text = ""

    def parse(self) -> dict:
        """
        解析招标文件
        """
        ext = os.path.splitext(self.tender_file)[1].lower()

        if ext == '.pdf':
            logger.info(f"Parsing PDF: {self.tender_file}")
            self.raw_text = PDFParser.parse(self.tender_file)
        elif ext in ['.docx', '.doc']:
            logger.info(f"Parsing DOCX: {self.tender_file}")
            self.raw_text = DOCXParser.parse(self.tender_file)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        logger.info(f"Extracted {len(self.raw_text)} characters from tender file")

        # 提取项目名称（从标题或第一行）
        project_name = self._extract_project_name()

        # 分类提取技术要求
        qualification_reqs = self._extract_by_keywords(self.QUALIFICATION_KEYWORDS)
        tech_reqs = self._extract_by_keywords(self.TECH_KEYWORDS)
        eval_reqs = self._extract_by_keywords(self.EVALUATION_KEYWORDS)

        return {
            'project_name': project_name,
            'raw_text': self.raw_text,
            'qualification_requirements': qualification_reqs,
            'technical_requirements': tech_reqs,
            'evaluation_criteria': eval_reqs
        }

    def _extract_project_name(self) -> str:
        """从文本中提取项目名称"""
        lines = self.raw_text.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 5 and len(line) < 200:
                # 排除明显不是项目名的行
                if not any(k in line for k in ['招标', '公告', '投标人', '授权']):
                    return line
        return "未识别项目名称"

    def _extract_by_keywords(self, keywords: list[str]) -> list[dict]:
        """
        根据关键词提取相关条款
        """
        clauses = []
        lines = self.raw_text.split('\n')

        current_clause = None
        clause_id = 0

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检查是否包含关键词
            if any(k in line for k in keywords):
                clause_id += 1
                current_clause = {
                    'id': f"clause_{clause_id}",
                    'title': self._clean_title(line),
                    'content': line,
                    'category': keywords[0] if keywords else 'unknown'
                }
                clauses.append(current_clause)
            elif current_clause and len(line) > 20:
                # 续写上一条款
                current_clause['content'] += " " + line

        return clauses

    def _clean_title(self, text: str) -> str:
        """清理标题，去除编号和特殊字符"""
        # 去除开头可能的编号如 "1.2.3" 或 "第一章"
        text = re.sub(r'^[\d一二三四五六七八九十]+[、\.．]+', '', text)
        text = re.sub(r'^[（）\(\)【】\[\]]+', '', text)
        return text.strip()[:100]  # 限制标题长度


# ============================================================
# 技术方案生成器
# ============================================================

class TechnicalProposalGenerator:
    """技术方案生成器"""

    def __init__(self, tender_file: str, output_file: str):
        self.tender_file = tender_file
        self.output_file = output_file
        self.requirements: list[Requirement] = []
        self.cases: list[ProjectCase] = []

    def parse_tender(self) -> dict:
        """
        解析招标文件，提取技术要求
        支持PDF和Word文件
        """
        logger.info(f"Parsing tender file: {self.tender_file}")
        parser = TenderParser(self.tender_file)
        return parser.parse()

    def extract_requirements(self, tender_data: dict) -> list[Requirement]:
        """
        从解析结果中提取技术要求，生成Requirement对象列表
        """
        requirements = []

        # 处理资格要求
        for clause in tender_data.get('qualification_requirements', []):
            req = Requirement(
                clause_id=clause['id'],
                title=clause['title'],
                content=clause['content'],
                response=self._generate_response(clause['title'], clause['content'], 'qualification'),
                chapter="第一章 资格证明"
            )
            requirements.append(req)

        # 处理技术要求
        for clause in tender_data.get('technical_requirements', []):
            req = Requirement(
                clause_id=clause['id'],
                title=clause['title'],
                content=clause['content'],
                response=self._generate_response(clause['title'], clause['content'], 'technical'),
                chapter="第二章 技术方案"
            )
            requirements.append(req)

        # 处理评审标准
        for clause in tender_data.get('evaluation_criteria', []):
            req = Requirement(
                clause_id=clause['id'],
                title=clause['title'],
                content=clause['content'],
                response=self._generate_response(clause['title'], clause['content'], 'evaluation'),
                chapter="第三章 评分标准响应"
            )
            requirements.append(req)

        logger.info(f"Extracted {len(requirements)} requirements from tender")
        return requirements

    def _generate_response(self, title: str, content: str, req_type: str) -> str:
        """
        根据条款类型生成我方响应
        """
        if req_type == 'qualification':
            return f"我方完全满足{title}要求，已准备好相关证明材料，可随时提供查验。"

        elif req_type == 'technical':
            return f"针对{title}，我方制定了详细的技术方案，确保满足所有技术指标要求。" \
                   f"具体实施方案包括：系统设计、设备选型、施工工艺、质量控制等全面措施。"

        elif req_type == 'evaluation':
            return f"我方将严格按照评分标准要求准备投标文件，确保在各项评审指标中获得高分。"

    def generate_chapter(self, chapter_title: str, requirements: list[Requirement]) -> str:
        """
        根据要求生成章节内容
        """
        content = f"## {chapter_title}\n\n"
        for req in requirements:
            content += f"### {req.title}\n\n"
            content += f"**招标文件要求**：{req.content}\n\n"
            content += f"**我方响应**：\n{req.response}\n\n"
        return content

    def generate_timeline(self, total_months: int, milestones: list[dict]) -> str:
        """
        生成项目实施计划（Markdown表格格式）
        """
        content = "### 项目实施计划\n\n"
        content += f"**总工期**：{total_months}个月\n\n"
        content += "| 阶段 | 里程碑 | 完成时间 | 交付物 |\n"
        content += "|------|--------|---------|-------|\n"

        for m in milestones:
            content += f"| {m['phase']} | {m['milestone']} | {m['time']} | {m['deliverable']} |\n"

        return content

    def generate_quality_measure(self) -> str:
        """
        生成质量保证措施章节
        """
        content = "## 质量保证措施\n\n"
        content += "### 质量管理体系\n\n"
        content += "我方已建立完善的质量管理体系，通过ISO9001:2015认证，涵盖：\n\n"
        content += "- 需求管理流程\n"
        content += "- 设计评审流程\n"
        content += "- 编码规范与代码审查\n"
        content += "- 测试管理流程\n"
        content += "- 变更管理流程\n\n"

        content += "### 质量控制节点\n\n"
        content += "| 阶段 | 质量控制点 | 检查标准 | 检查方法 |\n"
        content += "|------|---------|---------|---------|\n"
        content += "| 需求分析 | 需求评审 | 评审通过率100% | 同行评审 |\n"
        content += "| 系统设计 | 设计评审 | 评审通过率100% | 专家评审 |\n"
        content += "| 开发实现 | 代码审查 | 覆盖率≥80% | 自动化工具 |\n"
        content += "| 测试验收 | 验收测试 | 通过率≥95% | 测试报告 |\n\n"

        return content

    def generate_case_study(self, case: ProjectCase) -> str:
        """
        生成案例描述
        """
        content = f"**项目名称**：{case.name}\n\n"
        content += f"**客户**：{case.customer}\n\n"
        content += f"**合同金额**：¥{case.contract_amount:,.2f}元\n\n"
        content += f"**完成时间**：{case.completion_date}\n\n"
        content += f"**项目简介**：\n{case.description}\n\n"
        content += f"**技术亮点**：\n"
        for highlight in case.tech_highlights:
            content += f"- {highlight}\n"
        return content

    def build_proposal(self) -> str:
        """
        构建完整技术方案
        """
        tender_data = self.parse_tender()
        self.requirements = self.extract_requirements(tender_data)

        proposal = "# 技术投标书\n\n"
        proposal += f"**项目名称**：{tender_data.get('project_name', '[待填写]')}\n\n"
        proposal += f"**投标单位**：[待填写]\n\n"
        proposal += f"**日期**：{datetime.now().strftime('%Y年%m月%d日')}\n\n"
        proposal += "---\n\n"

        # 各章节
        for req in self.requirements:
            chapter = self.generate_chapter(req.chapter, [req])
            proposal += chapter

        proposal += self.generate_quality_measure()

        # 案例
        if self.cases:
            proposal += "## 典型案例\n\n"
            for case in self.cases:
                proposal += self.generate_case_study(case)
                proposal += "\n---\n\n"

        return proposal

    def save(self):
        """保存技术方案到文件"""
        proposal = self.build_proposal()
        with open(self.output_file, 'w', encoding='utf-8') as f:
            f.write(proposal)
        print(f"Technical proposal saved to: {self.output_file}")


# ============================================================
# 主程序
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="技术方案生成器")
    parser.add_argument("--tender-file", required=True, help="招标文件路径（PDF/Word）")
    parser.add_argument("--output", "-o", default="technical_proposal.md", help="输出文件路径")
    parser.add_argument("--cases-file", help="项目案例JSON文件")

    args = parser.parse_args()

    generator = TechnicalProposalGenerator(args.tender_file, args.output)

    if args.cases_file:
        with open(args.cases_file, 'r', encoding='utf-8') as f:
            cases_data = json.load(f)
            generator.cases = [ProjectCase(**c) for c in cases_data]

    generator.save()


if __name__ == "__main__":
    main()